from docx import Document
import re

# Step 1: Read and parse the voice-separated .txt file
def parse_conversation_file(file_path):
    attendees = set()
    discussion_points = []
    agenda_items = []
    start_time, end_time = None, None

    with open(file_path, 'r', encoding='utf-8') as file:
        for line in file:
            match = re.match(r'\[(\d{2}:\d{2}:\d{2})\] \[([^\]]+)\]: (.+)', line)
            if match:
                time_str, speaker, message = match.groups()
                attendees.add(speaker)

                # Extract general discussion points
                discussion_points.append(f"{speaker} discussed: {message.strip()}")

                # Capture start and end time
                time = time_str
                if not start_time:
                    start_time = time
                end_time = time

    # Simulated agenda extraction (customize based on your logic)
    if "Linear Equations" in " ".join(discussion_points):
        agenda_items.append("Linear Equations in Two Variables")

    return start_time, end_time, list(attendees), agenda_items, discussion_points

# Step 2: Generate MOM content dynamically
def generate_mom_content(start_time, end_time, attendees, agenda_items, discussion_points):
    mom_content = {
        "Meeting Information": {
            "Date": "[Not provided in transcript]",
            "Time": f"{start_time} to {end_time}" if start_time and end_time else "[Not provided in transcript]",
            "Location": "[Not provided in transcript]",
            "Attendees": ", ".join(attendees) if attendees else "[Not provided in transcript]"
        },
        "Agenda Items": agenda_items if agenda_items else ["[Not provided in transcript]"],
        "Discussion Points": discussion_points if discussion_points else ["[Not provided in transcript]"],
        "Decisions Made": "[Not provided in transcript]",
        "Action Items": "[Not provided in transcript]",
        "Next Meeting": "[Not provided in transcript]"
    }
    return mom_content

# Step 3: Create the MOM document
def create_mom_docx(mom_content, output_file):
    document = Document()
    document.add_heading('Meeting Minutes', level=1)

    # Add Meeting Information
    document.add_heading('Meeting Information', level=2)
    for key, value in mom_content["Meeting Information"].items():
        document.add_paragraph(f"{key}: {value}")

    # Add Agenda Items
    document.add_heading('Agenda Items', level=2)
    for item in mom_content["Agenda Items"]:
        document.add_paragraph(item, style='List Bullet')

    # Add Discussion Points
    document.add_heading('Discussion Points', level=2)
    for point in mom_content["Discussion Points"]:
        document.add_paragraph(point, style='List Number')

    # Add Decisions Made
    document.add_heading('Decisions Made', level=2)
    document.add_paragraph(mom_content["Decisions Made"])

    # Add Action Items
    document.add_heading('Action Items', level=2)
    document.add_paragraph(mom_content["Action Items"])

    # Add Next Meeting details
    document.add_heading('Next Meeting', level=2)
    document.add_paragraph(mom_content["Next Meeting"])

    # Save the document
    document.save(output_file)
    print(f"MOM saved successfully as {output_file}")

# Main execution
if __name__ == "__main__":
    # Input and output file paths
    conversation_file = '02-01-2025-15-34-05_transcription.txt'
    output_mom_file = 'Generated_MOM.docx'

    # Step 1: Parse the conversation file
    start_time, end_time, attendees, agenda_items, discussion_points = parse_conversation_file(conversation_file)

    # Step 2: Generate MOM content dynamically
    mom_content = generate_mom_content(start_time, end_time, attendees, agenda_items, discussion_points)

    # Step 3: Create and save the MOM document
    create_mom_docx(mom_content, output_mom_file)
